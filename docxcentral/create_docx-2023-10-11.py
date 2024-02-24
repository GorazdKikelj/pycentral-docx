from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxcompose.composer import Composer
from sortedcontainers import SortedList, SortedDict
import json
import base64
from pprint import pprint
import glob
from os.path import exists
from pycentral.base import ArubaCentralBase
from pycentral.configuration import Groups
from pprint import pprint
from docx2pdf import convert
from datetime import date, datetime
from pycentral.licensing import Subscriptions
from pycentral.device_inventory import Inventory

TEMPLATE_DOCX = "template/template.docx"
DIR_DOCX = "docx/"
DIR_BOM = "bom/"

# Create an instance of ArubaCentralBase using API access token
# or API Gateway credentials.
central_info = {
    "base_url": "https://apigw-eucentral3.central.arubanetworks.com",
    "token": {
        "access_token": "QknH9lX7HBBLaE4m2QIDitBW0PQmxE3C",
        "refresh_token": "xOdQfDyIKLqtyhd6NNBytxGztgXI2wz6",
    },
    "base_url": "https://apigw-uswest4.central.arubanetworks.com",
    # "customer_id": "58aeca167c4811ed957ab6cd43cbf282",
    # "client_id": "n5ultU0woilB3hMUc5dvCbctEDQyu2D6",
    # "client_secret": "ppdx3pRrk73iFqZJxPZNJiB4fYZzYO6x",
    # "username": "gorazd.kikelj@selectium.com",
    # "password": "Ruyagerakoc2055!",
}
token_store = {"type": "local", "path": "token"}
ssl_verify = True
central = ArubaCentralBase(
    central_info=central_info, token_store=token_store, ssl_verify=ssl_verify
)

visualrf_api = "/visualrf_api/v1/"


def get_central_data(apipath, apiparams={"offset": 0}):
    apiPath = apipath
    apiMethod = "GET"
    apiParams = apiparams or " "
    base_resp = central.command(
        apiMethod=apiMethod, apiPath=apiPath, apiParams=apiParams
    )
    return base_resp["msg"]


def get_campus_id() -> dict:
    return get_central_data(apipath=visualrf_api + "campus", apiparams={"offset": 0})


def get_buildings(campus_id) -> dict:
    return get_central_data(
        apipath=visualrf_api + f"campus/{campus_id}",
        apiparams={"offset": 0},
    )


def get_floors(building_id) -> dict:
    return get_central_data(
        apipath=visualrf_api + f"building/{building_id}",
        apiparams={"offset": 0, "units": "METERS"},
    )


def get_floor_data(floor_id) -> dict:
    return get_central_data(
        apipath=visualrf_api + f"floor/{floor_id}/access_point_location",
        apiparams={"offset": 0, "units": "METERS"},
    )


def save_floorplans() -> dict:
    campuses = get_campus_id()
    floor_dict = {}
    buildings = get_buildings(campus_id=campuses["campus"][0]["campus_id"])
    for building in buildings["buildings"]:
        floors = get_floors(building_id=building["building_id"])
        for floor in floors["floors"]:
            floor_dict[floor["floor_name"]] = {
                "floor_id": floor["floor_id"],
                "ap": {},
            }
            img = base64.b64decode(
                get_central_data(
                    apipath=f"/visualrf_api/v1/floor/{floor['floor_id']}/image",
                    apiparams={"offset": 0},
                ),
                validate=True,
            )
            with open(
                f"images/{building['building_name']}_floor_{floor['floor_level']}.png",
                "wb",
            ) as f:
                f.write(img)

            floor_data = get_floor_data(floor_id=floor["floor_id"])
            for ap in floor_data["access_points"]:
                floor_dict[floor["floor_name"]]["ap"].update(
                    {ap["ap_name"]: ap["ap_id"]}
                )
    return floor_dict


def get_rf_groups(group_name) -> list:
    apipath = f"/configuration/v1/dot11a_radio_profiles/{group_name}"
    return get_central_data(apipath=apipath)


def get_central_groups() -> list:
    g = Groups()
    group_list = g.get_groups(central)["msg"]

    return group_list.get("data")


def get_wlan_list(group_name) -> dict:
    apipath = f"/configuration/full_wlan/{group_name}"
    data = get_central_data(apipath=apipath)
    if type(data) is dict:
        pprint(data)
        return data.get("description")
    return json.loads(data)


def set_column_width(column, width) -> None:
    for cell in column.cells:
        cell.width = Cm(width)


def get_per_ap_settings(serial_no) -> list:
    #    apimethod = "GET"
    apipath = f"/configuration/v1/ap_settings_cli/{serial_no}"
    return get_central_data(apipath=apipath)


def sort_ap_list(file_mask, reverse=False) -> list:
    file_list = []
    for picture_file in glob.glob(file_mask):
        file_list.append(picture_file)
    file_list.sort(reverse=reverse)
    return file_list


def sort_ap_dict(ap_list: dict) -> list[list]:
    new_ap = []
    for ap_item in ap_list:
        k = [ap_item, ap_list[ap_item]["name"]]
        new_ap.append(k)
    return sorted(new_ap, key=lambda x: x[1])


def add_ap_row(table, label, value, offset=0):
    row_cells = table.add_row().cells
    row_cells[0].text = f"{label}:"
    row_cells[0].paragraphs[0].style = "Table Rowhead 8 pt"
    row_cells[0].width = Cm(2)
    row_cells[1].text = f"{value}"
    row_cells[1].paragraphs[0].style = "Table Body 8pt"
    return row_cells


def add_picture_to_cell(paragraph, picture, width=None, height=None) -> None:
    run = paragraph.add_run()
    try:
        run.add_picture(picture, width=width, height=height)
    except FileNotFoundError:
        pass
    return None


def sort_list(data, key) -> list:
    """
    Sort list of dictionaries by a key "name"

    Return list of APs sorted by name and position index of the
    data inside the data list in tupple (name, index).

    """
    name_list = []
    idx = 0
    for name in data:
        name_list.append((name[key], idx))
        idx += 1
    return SortedList(name_list)


def add_ap_to_page(document, item) -> None:
    print(f"Add AP {item['name']} {item['serial']}")
    ap_data = get_per_ap_settings(serial_no=item["serial"])
    document.add_paragraph(
        f'AP: {item["name"]}', style="Aruba body Quote text 2 Orange Arial 16pt"
    )
    main_table = document.add_table(rows=1, cols=2)
    hdr_row = main_table.rows[0].cells
    table = hdr_row[0].add_table(rows=1, cols=2)
    table.autofit = True
    table.style = "Table Grid"
    table.columns[0].width = Cm(3)
    table.cell(0, 0).text = "Parameter"
    table.cell(0, 0).width = Cm(3)
    table.cell(0, 1).text = "Value"
    row_cells = add_ap_row(table=table, label="Site", value=item["site"])
    row_cells = add_ap_row(table=table, label="AP Group", value=item["group_name"])
    row_cells = add_ap_row(table=table, label="AP Model", value=item["model"])
    row_cells = add_ap_row(table=table, label="Serial No", value=item["serial"])
    row_cells = add_ap_row(table=table, label="Labels", value=item["labels"])
    row_cells = add_ap_row(table=table, label="MAC address", value=item["macaddr"])
    row_cells = add_ap_row(table=table, label="Mesh role", value=item["mesh_role"])
    row_cells = add_ap_row(
        table=table, label="IP address", value=item.get("ip_address")
    )
    row_cells = add_ap_row(
        table=table, label="Public IP address", value=item.get("public_ip_address")
    )
    row_cells = add_ap_row(
        table=table, label="Firmware version", value=item.get("firmware_version")
    )
    row_cells = add_ap_row(
        table=table, label="SSID count", value=item.get("ssid_count")
    )
    try:
        rf_zone = ap_data[11]
    except IndexError:
        rf_zone = "default default"
    row_cells = add_ap_row(table=table, label="RF zone", value=rf_zone.split()[1])

    radio_table = hdr_row[1].add_table(rows=1, cols=2)
    radio_table.autofit = False
    radio_table.style = "Table Grid"
    radio_table.columns[0].width = Cm(2)
    radio_table.cell(0, 0).text = "Radio"
    radio_table.cell(0, 0).width = Cm(2)
    radio_table.cell(0, 1).text = "Parameter"
    for radios in item["radios"]:
        row_cells = add_ap_row(
            table=radio_table, label=f'Radio {radios["index"]}', value=""
        )

        p = row_cells[1].paragraphs[0].clear()
        table_cell = row_cells[1].add_table(rows=0, cols=2)
        for radio in radios:
            if radio in [
                "macaddr",
                "radio_name",
                "radio_type",
                "spatial_stream",
                "tx_power",
            ]:
                table_cells = add_ap_row(
                    table=table_cell,
                    label=radio,
                    value=str(radios.get(radio)),
                    offset=0,
                )

    p = document.add_paragraph("Notes: ", style="Table Rowhead 8 pt")
    p = document.add_paragraph(item.get("notes"), style="Table Body 8pt")
    p = document.add_paragraph("Location: ", style="Table Rowhead 8 pt")
    for picture_file in glob.glob(f"images/{item['serial']}-*_location.png"):
        p = document.add_paragraph(" ")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(picture_file, width=Cm(13))

    p = document.add_paragraph(" ", style="Table Rowhead 8 pt")

    idx = 4
    for picture_file in sort_ap_list(f"images/{item['serial']}*.jpg"):
        if idx > 2:
            p = document.add_paragraph("\n")
            r = p.add_run()
            idx = 1
        r.add_picture(picture_file, width=Cm(7.5))
        r.add_text(" ")
        idx += 1

    ap_config = get_central_data(
        apipath=f"/configuration/v1/devices/{item['serial']}/configuration",
        apiparams={"limit": 0},
    )

    p = document.add_page_break()
    p = document.add_paragraph()
    p.add_run("Configuration")
    p.style = "Table Rowhead 8 pt"
    p = document.add_paragraph()
    p.add_run(ap_config).font.name = "Consolas"
    p.style = "Table Body 8pt"

    return None


def add_document_header(document, item) -> None:
    document.add_paragraph(item, style="Aruba Cover: Main title Orange Arial Bold 36pt")
    document.add_paragraph(
        f"\n\nMestna občina Celje\n\nKonfiguracija {date.today()}\n\n\n\n",
        style="Aruba Cover: Subheading CAPS Dark Blue Arial 20pt",
    )

    return None


def add_site_document(item, ap_list, data) -> None:
    document = Document(TEMPLATE_DOCX)
    add_document_header(document=document, item=item["site_name"])
    document.add_paragraph(
        f'Site: {item["site_name"]}',
        style="Aruba body Quote text 2 Orange Arial 16pt",
    )
    table = document.add_table(rows=1, cols=2)
    table.autofit = True
    table.style = "Table Grid"
    table.columns[0].width = Cm(3)
    table.cell(0, 0).width = Cm(3)
    row_cells = add_ap_row(
        table=table,
        label="Number of devices",
        value=str(item["associated_device_count"]),
    )
    row_cells = add_ap_row(table=table, label="Address", value=item.get("address"))
    row_cells = add_ap_row(table=table, label="Post Code", value=item.get("zipcode"))
    row_cells = add_ap_row(table=table, label="City", value=item.get("city"))
    row_cells = add_ap_row(table=table, label="Country", value=item.get("country"))
    row_cells = add_ap_row(table=table, label="Longitude", value=item.get("longitude"))
    row_cells = add_ap_row(table=table, label="Latitude", value=item.get("latitude"))
    for ap in ap_list[item["site_name"]]:
        row_cells = add_ap_row(table=table, label="AP", value=data["aps"][ap]["name"])

    document.add_page_break()
    for ap in ap_list[item["site_name"]]:
        add_ap_to_page(document=document, item=data["aps"][ap])
        document.add_page_break()

    doc_filename = f"{item['site_name']}.docx"
    #    filename = f"{DIR_BOM}{doc_filename[4:len(doc_filename)]}"
    filename = f"{DIR_BOM}{doc_filename}"
    print(f"BOM FILENAME : {doc_filename}....{filename}")
    add_bom_list(
        document=document, doc_filename=f"{DIR_DOCX}{doc_filename}", filename=filename
    )
    convert_docx_to_pdf(doc_filename=f"{DIR_DOCX}{doc_filename}")

    return None


def add_rf_group_to_page(group_name) -> None:
    """
    Create document with all RF group data for AP Group

    """
    document = Document(TEMPLATE_DOCX)
    data = get_rf_groups(group_name=group_name)
    add_document_header(document=document, item=f"Configuration group\n{group_name}")
    document.add_page_break()
    for groups in data:
        document.add_paragraph(
            f"RF Group: {group_name}",
            style="Aruba body Quote text 2 Orange Arial 16pt",
        )
        table = document.add_table(rows=1, cols=2)
        table.autofit = True
        table.style = "Table Grid"

        for group in groups:
            try:
                row_cells = add_ap_row(table=table, label=group, value=groups[group])
            except TypeError:
                pass

        document.add_page_break()

    doc_filename = f"{DIR_DOCX}{group_name}_rf_groups.docx"
    document.save(doc_filename)
    convert_docx_to_pdf(doc_filename=doc_filename)
    return None


def add_wlan_group_to_page(group_name) -> None:
    """
    Create document with all RF group data for AP Group

    """

    document = Document(TEMPLATE_DOCX)
    data = get_wlan_list(group_name=group_name)
    if type(data) is not dict:
        print(f"No data returned for WLANs on group {group_name}")
        return None

    add_document_header(document=document, item=f"Configuration group\n{group_name}")
    document.add_page_break()

    for groups in data.get("wlans"):
        document.add_paragraph(
            f"WLAN: {groups['name']}",
            style="Aruba body Quote text 2 Orange Arial 16pt",
        )
        table = document.add_table(rows=1, cols=2)
        table.autofit = True
        table.style = "Table Grid"

        for group in groups:
            tmp_value = groups[group]
            value = tmp_value if not isinstance(tmp_value, dict) else tmp_value["value"]
            row_cells = add_ap_row(table=table, label=group, value=value)
        document.add_page_break()
    doc_filename = f"{DIR_DOCX}{group_name}_wlan_groups.docx"
    document.save(doc_filename)
    convert_docx_to_pdf(doc_filename=doc_filename)

    return None


def add_bom_list(document, doc_filename, filename) -> None:
    """Add additional document to the master document file"""
    comp = Composer(document)
    if exists(filename):
        sub_doc = Document(filename)
        comp.append(sub_doc)
    else:
        print(f"ERROR: BOM file does not exist. {filename}")

    comp.save(doc_filename)
    return None


def convert_docx_to_pdf(doc_filename) -> None:
    if not exists(doc_filename):
        return None
    print(f"Convert docx to pdf {doc_filename}")
    convert(doc_filename)

    return None


def add_subscription_keys() -> None:
    subs = Subscriptions()
    data_raw = subs.get_user_subscription_keys(central)
    data_msg = data_raw.get("msg")
    if data_msg.get("status") != 200:
        data = {}
        print(data_msg)
        return None

    data = data_msg.get("subscriptions")
    document = Document(TEMPLATE_DOCX)
    add_document_header(document=document, item=f"Subscriptions")
    document.add_page_break()

    document.add_paragraph(
        f"Subscriptions",
        style="Aruba body Quote text 2 Orange Arial 16pt",
    )
    table = document.add_table(rows=1, cols=8)
    table.autofit = True
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "SKU"
    hdr_cells[0].paragraphs[0].style = "Table Rowhead 8 pt"
    hdr_cells[1].text = "License Type"
    hdr_cells[1].paragraphs[0].style = "Table Rowhead 8 pt"
    hdr_cells[2].text = "Quantity"
    hdr_cells[2].paragraphs[0].style = "Table Rowhead 8 pt"
    hdr_cells[3].text = "Available"
    hdr_cells[3].paragraphs[0].style = "Table Rowhead 8 pt"
    hdr_cells[4].text = "Active"
    hdr_cells[4].paragraphs[0].style = "Table Rowhead 8 pt"
    hdr_cells[5].text = "Subscription Key"
    hdr_cells[5].paragraphs[0].style = "Table Rowhead 8 pt"
    hdr_cells[6].text = "Start Date"
    hdr_cells[6].paragraphs[0].style = "Table Rowhead 8 pt"
    hdr_cells[7].text = "End Date"
    hdr_cells[7].paragraphs[0].style = "Table Rowhead 8 pt"

    for subscription in data:
        if "EVAL" in subscription.get("sku"):
            continue
        row_cells = table.add_row().cells
        row_cells[0].text = subscription.get("sku")
        row_cells[0].paragraphs[0].style = "Table Body 8pt"
        row_cells[1].text = subscription.get("license_type")
        row_cells[1].paragraphs[0].style = "Table Body 8pt"
        row_cells[2].text = f'{subscription.get("quantity")}'
        row_cells[2].paragraphs[0].style = "Table Body 8pt"
        row_cells[3].text = f'{subscription.get("available")}'
        row_cells[3].paragraphs[0].style = "Table Body 8pt"
        row_cells[4].text = f'{subscription.get("active")}'
        row_cells[4].paragraphs[0].style = "Table Body 8pt"
        row_cells[5].text = f'{subscription.get("subscription_key")}'
        row_cells[5].paragraphs[0].style = "Table Body 8pt"
        row_cells[
            6
        ].text = f'{datetime.fromtimestamp(int(subscription.get("start_date")) / 1000).strftime("%d.%m.%Y")}'
        row_cells[6].paragraphs[0].style = "Table Body 8pt"
        row_cells[
            7
        ].text = f'{datetime.fromtimestamp(int(subscription.get("end_date")) / 1000  ).strftime("%d.%m.%Y")}'
        row_cells[7].paragraphs[0].style = "Table Body 8pt"

    doc_filename = f"{DIR_DOCX}subscriptions.docx"
    document.save(doc_filename)
    convert_docx_to_pdf(doc_filename=doc_filename)
    return None


def add_device_inventory(ap_list: dict) -> None:
    inve = Inventory()
    msg = inve.get_inventory(central, limit=120)["msg"]
    data = msg["devices"]
    total = msg["total"]
    for ap in data:
        sn = ap["serial"]
        if sn not in ap_list:
            ap_list[sn] = {"name": "", "site": ""}
        ap_list[sn]["aruba_part_no"] = ap.get("aruba_part_no")
        ap_list[sn]["device_type"] = ap.get("device_type")
        ap_list[sn]["macaddr"] = ap.get("macaddr")
        ap_list[sn]["model"] = ap.get("model")
        ap_list[sn]["serial"] = ap.get("serial")
        ap_list[sn]["subscription_key"] = ap.get("subscription_key")
        ap_list[sn]["tier_type"] = ap.get("tier_type")

    document = Document(TEMPLATE_DOCX)
    add_document_header(document=document, item=f"Device Inventory")
    document.add_page_break()

    document.add_paragraph(
        f"Subscriptions",
        style="Aruba body Quote text 2 Orange Arial 16pt",
    )
    document.add_paragraph(
        f"Total devices: {total}",
        style="Aruba body Quote text 2 Orange Arial 16pt",
    )

    table = document.add_table(rows=1, cols=7)
    table.autofit = True
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Part No"
    hdr_cells[0].paragraphs[0].style = "Table Rowhead 8 pt"
    hdr_cells[1].text = "Device type"
    hdr_cells[1].paragraphs[0].style = "Table Rowhead 8 pt"
    hdr_cells[2].text = "MAC Address"
    hdr_cells[2].paragraphs[0].style = "Table Rowhead 8 pt"
    hdr_cells[3].text = "Model"
    hdr_cells[3].paragraphs[0].style = "Table Rowhead 8 pt"
    hdr_cells[4].text = "Serial"
    hdr_cells[4].paragraphs[0].style = "Table Rowhead 8 pt"
    hdr_cells[5].text = "Subscription Key"
    hdr_cells[5].paragraphs[0].style = "Table Rowhead 8 pt"
    hdr_cells[6].text = "Tier"
    hdr_cells[6].paragraphs[0].style = "Table Rowhead 8 pt"

    for key, value in sort_ap_dict(ap_list):
        #    for subscription in data:
        subscription = ap_list[key]
        row_cells = table.add_row().cells
        row = table.rows[-1]
        a, b, d, c = row.cells[0:4]
        c = a.merge(c)
        try:
            c.text = f'\n{subscription.get("name")}'
        except KeyError:
            c.text = "\n"
        c.paragraphs[0].style = "Table Body 8pt"

        a, b, d = row.cells[4:7]
        c = a.merge(d)
        try:
            c.text = f'\n{subscription.get("site")}'
        except KeyError:
            c.text = "\n"
        c.paragraphs[0].style = "Table Body 8pt"

        row_cells = table.add_row().cells
        row_cells[0].text = subscription.get("aruba_part_no")
        row_cells[0].paragraphs[0].style = "Table Body 8pt"
        row_cells[1].text = subscription.get("device_type")
        row_cells[1].paragraphs[0].style = "Table Body 8pt"
        row_cells[2].text = f'{subscription.get("macaddr")}'
        row_cells[2].paragraphs[0].style = "Table Body 8pt"
        row_cells[3].text = f'{subscription.get("model")}'
        row_cells[3].paragraphs[0].style = "Table Body 8pt"
        row_cells[4].text = f'{subscription.get("serial")}'
        row_cells[4].paragraphs[0].style = "Table Body 8pt"
        row_cells[5].text = f'{subscription.get("subscription_key")}'
        row_cells[5].paragraphs[0].style = "Table Body 8pt"
        row_cells[6].text = f'{subscription.get("tier_type")}'
        row_cells[6].paragraphs[0].style = "Table Body 8pt"

    doc_filename = f"{DIR_DOCX}device_inventory.docx"
    document.save(doc_filename)
    convert_docx_to_pdf(doc_filename=doc_filename)
    return None


def add_sites_to_page(site_name, sites) -> dict:
    apiparams = {
        #        "group": group_name,
        "site": site_name,
        "offset": 0,
        "calculate_totals": "true",
        "calculate_client_count": "true",
        "calculate_ssid_count": "true",
        "show_resource_details": "true",
    }
    apipath = "/monitoring/v2/aps"

    data = get_central_data(apipath=apipath, apiparams=apiparams)
    #    print(f"--AP Data ---")
    #    pprint(data)

    """
Pripravi seznam APjev sortiran po imenu in doda APje na Site.
    """
    ap_list = {}
    ap_list_by_sn = {}
    for name, idx in sort_list(data=data["aps"], key="name"):
        item = data["aps"][idx]
        ap_list_by_sn[item["serial"]] = {
            "name": item["name"],
            "site": item["site"],
            "group": item["group_name"],
        }
        if item["site"] not in ap_list.keys():
            ap_list[item["site"]] = [idx]
        else:
            ap_list[item["site"]].append(idx)

    print(f"-- AP List -- len = {len(ap_list)}")

    """ Sites data """
    for name, idx in sort_list(data=sites["sites"], key="site_name"):
        item = sites["sites"][idx]
        if item["site_name"] not in ap_list.keys():
            print(f'Ignored site: {item["site_name"]}')
        else:
            print(f'Processed site: {item["site_name"]}')
            add_site_document(item=item, ap_list=ap_list, data=data)

    return ap_list_by_sn


"""
Podatki o posamezni lokaciji (site)

 /visualrf_api/v1/campus --> campus list
 
 /visualrf_api/v1/campus/<campus_id> --> seznam lokacij (sites)
 
/visualrf_api/v1/building/<building_id> --> podatki o site-u in floor planih

/visualrf_api/v1/floor/<floor_id>/access_point_location --> podatki o floor planu in lokacijah APjev na floor planu

/visualrf_api/v1/floor/<floor_id>/image --> background slika floor plana

"""


if __name__ == "__main__":
    add_subscription_keys()
    all_groups = get_central_groups()
    sites = get_central_data(
        apipath="/central/v2/sites", apiparams={"calculate_totals": "true"}
    )
    # pprint(sites)
    # all_groups = [["Tehnopark"]]
    ap_list = {}
    for group in all_groups:
        print(
            f"Working on group {group} ------------------------------------------------"
        )
        add_rf_group_to_page(group_name=group[0])
        add_wlan_group_to_page(group_name=group[0])

    for site in sites["sites"]:
        ap_list = ap_list | add_sites_to_page(site_name=site["site_name"], sites=sites)
        print(
            f"Adding APs to list site {site['site_name']} ap list size {len(ap_list)}"
        )
    add_device_inventory(ap_list=ap_list)

    exit()
